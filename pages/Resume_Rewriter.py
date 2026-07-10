import streamlit as st
import io
from docx import Document
from models.recommendation import suggest_bullet_rewrites, generate_optimized_summary

def show():
    # Load custom styles
    with open("assets/custom.css") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
        
    st.markdown("<h2 class='gradient-text'>Resume Rewriter Workshop</h2>", unsafe_allow_html=True)
    st.write("Optimize your resume's language and sections to match the target job description.")
    
    if "sacs_result" not in st.session_state:
        st.warning("⚠️ No active analysis found. Please run the analysis on the 'Resume Analysis' page first.")
        return
        
    results = st.session_state.sacs_result
    skills_report = results["skills_report"]
    matched_skills = skills_report["matched_tech"] + skills_report["matched_soft"]
    
    tab1, tab2 = st.tabs(["✍️ Bullet Point Improver", "📄 Auto-Generated Sections"])
    
    with tab1:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.subheader("Interactive Bullet Improver")
        st.write("Paste a weak or passive bullet point from your experience list, and our AI rewriter will formulate strong, active, metric-focused versions.")
        
        weak_input = st.text_input(
            "Enter passive description (e.g., 'I was responsible for maintaining the database')",
            value="I was responsible for maintaining the database and helping with coding"
        )
        
        if weak_input:
            st.write("")
            suggestions = suggest_bullet_rewrites(weak_input)
            
            st.markdown("##### 🚀 Professional Recommendations:")
            for idx, sug in enumerate(suggestions, 1):
                st.markdown(f"<div class='info-callout'><b>Option {idx}:</b> {sug}</div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)
        
    with tab2:
        st.markdown("<div class='glass-card'>", unsafe_allow_html=True)
        st.subheader("Optimized Resume Content")
        st.write("Here are tailored resume sections based on your match analysis. Review and copy these directly or download the DOCX file below.")
        
        # 1. Professional Summary
        st.markdown("##### 1. Optimized Professional Summary")
        existing_summary = results["parsed_resume"]["sections"].get("summary", "")
        opt_summary = generate_optimized_summary(existing_summary, results["parsed_jd"]["raw_text"], matched_skills)
        summary_edit = st.text_area("Tailored Summary", value=opt_summary, height=120)
        
        # 2. Optimized Skills Section
        st.markdown("##### 2. ATS-Optimized Skills Block")
        st.write("Grouped list of your matching skills categorized for high ATS readability:")
        
        # Build skills text block
        skills_text_parts = []
        # Group matched skills
        tech_matched = skills_report["matched_tech"]
        transfer_matched = skills_report["transfer_tech"]
        soft_matched = skills_report["matched_soft"]
        
        if tech_matched:
            skills_text_parts.append(f"Technical Core: {', '.join(tech_matched)}")
        if transfer_matched:
            skills_text_parts.append(f"Additional Competencies: {', '.join(transfer_matched)}")
        if soft_matched:
            skills_text_parts.append(f"Professional Methodologies: {', '.join(soft_matched)}")
            
        skills_text_block = "\n".join(skills_text_parts)
        skills_edit = st.text_area("Grouped Skills", value=skills_text_block, height=100)
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Word document downloader
        st.write("")
        
        # Generate DOCX in memory
        doc = Document()
        doc.add_heading("Optimized Resume Sections", level=1)
        
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary_edit)
        
        doc.add_heading("Technical & Professional Skills", level=2)
        doc.add_paragraph(skills_edit)
        
        doc.add_heading("Suggested Project/Experience Bullets", level=2)
        doc.add_paragraph("Replace passive verbs in your experience with active ones. Example improvements:")
        if weak_input:
            for s in suggestions:
                doc.add_paragraph(s, style='List Bullet')
                
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        st.download_button(
            label="Download Polished Resume Sections (.docx) 📥",
            data=docx_buffer,
            file_name="optimized_resume_sections.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
